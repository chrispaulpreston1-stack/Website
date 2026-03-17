"""Generate Word document: Email to Watershedd (Marzouk Al-Bader) re Bobtails, Mawgan Porth."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)
    return h

def add_body(text):
    return doc.add_paragraph(text)

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('\u2500' * 70)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Subject line
p = doc.add_paragraph()
run = p.add_run('Subject: ')
run.bold = True
run.font.size = Pt(12)
run = p.add_run('Bobtails, Mawgan Porth \u2014 Fascinating Project (Academic Interest)')
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(16)

add_separator()

# Opening
add_body('Hi Marzouk,')

add_body(
    'My name\u2019s Chris from PF & Co Site Intelligence. I hope you don\u2019t mind the cold email '
    '\u2014 your Bobtails project at Trenance came across us via a LinkedIn post by Sara de Barros '
    'discussing the recent landscaping application, and we ended up going down quite the rabbit hole.'
)

add_body(
    'We\u2019re building an AI-powered planning intelligence system and we\u2019ve been using '
    'real-world applications as learning exercises to stress-test our constraint detection and QA '
    'processes. Bobtails turned out to be one of the most interesting case studies we\u2019ve looked at.'
)

add_separator()

# --- The Property History ---
add_heading_styled('The Property History', level=2)

add_body(
    'Before anything else \u2014 what a site. We dug into the backstory and it\u2019s a genuinely '
    'compelling piece of Cornwall\u2019s recent architectural history. The journey from original cottage '
    'through to the replacement eco-home you\u2019ve designed is exactly the kind of sensitive coastal '
    'work that\u2019s hard to get right. The approach of working with the landscape rather than against '
    'it \u2014 the living roofs, locally sourced Cornish stone and timber, the low-energy design '
    'philosophy \u2014 is clear even from the outside.'
)

add_body(
    'We also noticed that several newspapers have attributed a connection to Sir Arthur Conan Doyle '
    'and Tom Baker to the Mawgan Porth cottage. We actually traced this back and it appears to be a '
    'case of journalists mixing up two properties \u2014 the Conan Doyle and Tom Baker stories relate '
    'to Highwell House in Crowborough, East Sussex, not Bobtails. The Conan Doyle Encyclopedia lists '
    'no Cornwall residences. Bit of an unexpected detour but quite satisfying to untangle!'
)

add_separator()

# --- The Design ---
add_heading_styled('The Design', level=2)

add_body(
    'The replacement dwelling is a genuinely impressive piece of work. Achieving planning consent for '
    'a contemporary eco-home on an exposed clifftop site within the Lanherne and Watergate AGLV, '
    'adjacent to the South West Coast Path, and with the level of public interest that Mawgan Porth '
    'attracts \u2014 that\u2019s not straightforward. The fact that it secured approval through '
    'PA21/12699 with what was clearly a thorough pre-application process speaks to the quality of the '
    'design response.'
)

add_separator()

# --- System Flags ---
add_heading_styled('A Few Things Our System Flagged', level=2)

add_body(
    'When we ran the site and the latest landscaping application through our system purely as a '
    'learning exercise, a handful of things came back that we found genuinely interesting. Not as '
    'criticisms at all \u2014 more that they raised questions we\u2019d love to understand the '
    'thinking behind.'
)

# Native planting
add_heading_styled('Native Planting Ratio', level=3)
add_body(
    'Our system flagged that the planting schedule appears to be predominantly non-native (Olearia '
    'traversii, Ampelodesmos mauritanicus, Muehlenbeckia, Phillyrea latifolia, Lonicera '
    'alseuosmoides) with sea thrift as the main confirmed native species. Cornwall\u2019s Climate '
    'Emergency DPD sets a benchmark of at least 50% pollinator-friendly planting of predominantly '
    'native species. We\u2019re curious whether this has come up in discussions with the council, or '
    'whether there\u2019s a design rationale that addresses it \u2014 perhaps the coastal exposure '
    'demands the hardier non-native species, or the DPD requirement is being interpreted differently '
    'for garden landscaping versus new-build schemes?'
)

# Muehlenbeckia
add_heading_styled('Muehlenbeckia Near the CWS', level=3)
add_body(
    'This was the one that really interested us. BSBI data shows Muehlenbeckia complexa naturalising '
    'and expanding its range in south-western England, specifically colonising cliff and dune habitats. '
    'The Mawgan Porth to Newquay County Wildlife Site (maritime cliff and slope, Section 41 Priority '
    'Habitat) is approximately 60 metres from the property. It\u2019s not Schedule 9 listed so '
    'there\u2019s no legal issue, but we wondered whether the ecology officer has raised any questions '
    'about the risk of spread into the adjacent designated habitat, or whether the contained nature of '
    'the planting (the \u201cevergreen curtain\u201d along the bank) is considered sufficient mitigation?'
)

# Corten steel
add_heading_styled('Corten Steel in the Coastal Environment', level=3)
add_body(
    'The Corten steps are a beautiful design choice and we completely understand the aesthetic '
    'rationale \u2014 the weathered finish sits naturally in a coastal landscape. Our system flagged '
    'a durability note though: Corten relies on a wet-dry cycle to form its protective oxide patina, '
    'and persistent chloride-rich salt spray within a few hundred metres of the coast can prevent that '
    'stabilising, leading to ongoing pitting rather than the intended even weathering. Has this been a '
    'consideration in the specification, or is there a marine-grade treatment being applied? Just '
    'curious \u2014 the Cornish granite you\u2019ve used elsewhere on the project is obviously '
    'bombproof.'
)

# AGLV and SWCP
add_heading_styled('AGLV and SWCP Context', level=3)
add_body(
    'We noted the site sits within the Lanherne and Watergate AGLV (saved Policy 14, Restormel Local '
    'Plan \u2014 confirmed as carrying real planning weight by the Court of Appeal in Corbett v '
    'Cornwall Council [2020] EWCA Civ 508) and that the South West Coast Path runs immediately along '
    'the cliff edge at Trenance. Has the council asked for any landscape and visual impact assessment '
    'for the landscaping works, given these designations? Or is the view that softening landscaping is '
    'inherently beneficial to the AGLV and doesn\u2019t need a formal assessment?'
)

add_separator()

# --- Purely Academic ---
add_heading_styled('Purely Academic', level=2)

add_body(
    'I want to stress \u2014 this is entirely a learning exercise for us. We\u2019re not involved in '
    'the area, we have no interest in the application, and we think the overall design approach is '
    'excellent. The permeable grass driveway with Greenstones pavers, the Cornish boulders, the intent '
    'to soften the building into the landscape rather than intensify it \u2014 that\u2019s exactly the '
    'kind of thoughtful, site-specific response that planning officers and landscape consultants want '
    'to see.'
)

add_body(
    'We\u2019re just genuinely interested in how these kinds of edge-case technical questions play out '
    'in practice with Cornwall Council, because it helps us calibrate our system against real-world '
    'outcomes rather than just policy text.'
)

add_body(
    'If you ever have a spare ten minutes and fancy a chat about any of the above, we\u2019d be '
    'delighted. And if you\u2019d like to see what our system does \u2014 happy to run any site '
    'through it, on the house.'
)

add_body('Fantastic work on Bobtails. Looking forward to seeing how the landscaping shapes up.')

# Sign off
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Best regards,')
p = doc.add_paragraph()
run = p.add_run('Chris')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('PF & Co Site Intelligence')
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

# --- Save ---
output_path = os.path.join(
    r'C:\Users\chris\pf-&-co-construction  - web\docs',
    'Email-Watershedd-Bobtails-Mawgan-Porth.docx'
)
doc.save(output_path)
print(f'Document saved to: {output_path}')
