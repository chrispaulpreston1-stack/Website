"""Generate Word document: Email to Dan Mitchell re Camborne scheme."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_body(text):
    return doc.add_paragraph(text)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 70)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)

def add_table_row_data(table, row_idx, cells_data, bold_first=False):
    row = table.rows[row_idx]
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if bold_first and i == 0:
            run.bold = True


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Subject line
p = doc.add_paragraph()
run = p.add_run('Subject: ')
run.bold = True
run.font.size = Pt(12)
run = p.add_run('Your Camborne Apartment Scheme — We Used It as a Case Study (Impressive Work)')
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(16)

add_separator()

# Opening
add_body('Hi Dan,')

add_body(
    'My name\u2019s Chris from PF & Co Site Intelligence. I came across your LinkedIn post about the '
    'Camborne apartment scheme \u2014 the 10-unit refusal turned into a 9-unit approval with full '
    'officer and Town Council support.'
)

add_body(
    'I hope you don\u2019t mind, but we ran the site through our planning intelligence system as a '
    'learning exercise. I wanted to share what our system pulled back and ask whether our understanding '
    'is right, because if it is, it\u2019s a genuinely smart piece of work on your part.'
)

add_separator()

# --- SECTION 1: Site Identification ---
add_heading_styled('Our Site Identification', level=2)

add_body(
    'Our constraint detection system points to the former Royal British Legion site, Gurneys Lane, '
    'Camborne, TR14 8JP (or the immediately adjacent Gas Street area). The evidence that flagged this was:'
)

add_bullet('', bold_prefix='Registered brownfield site ')
doc.paragraphs[-1].clear()
p = doc.paragraphs[-1]
run = p.add_run('Registered brownfield site')
run.bold = True
p.add_run(' (BR165 on planning.data.gov.uk), 0.05 hectares')

add_bullet('', bold_prefix='Original brownfield register description: ')
doc.paragraphs[-1].clear()
p = doc.paragraphs[-1]
run = p.add_run('Original brownfield register description: ')
run.bold = True
run2 = p.add_run('"Demolition and Construction of 10 new flats"')
run2.italic = True
p.add_run(' \u2014 matches your post exactly')

add_bullet(
    '', bold_prefix='Camborne Town Centre Conservation Area '
)
doc.paragraphs[-1].clear()
p = doc.paragraphs[-1]
run = p.add_run('Camborne Town Centre Conservation Area')
run.bold = True
p.add_run(' (designated September 2004)')

add_bullet('', bold_prefix='Area A5 ')
doc.paragraphs[-1].clear()
p = doc.paragraphs[-1]
run = p.add_run('Area A5 (Camborne & Redruth with Portreath)')
run.bold = True
p.add_run(
    ' of the Cornwall and West Devon Mining Landscape World Heritage Site \u2014 the area with '
    'the highest concentration of historic mining sites anywhere in the world'
)

add_bullet('', bold_prefix='Previous planning permission ')
doc.paragraphs[-1].clear()
p = doc.paragraphs[-1]
run = p.add_run('Previous planning permission')
run.bold = True
p.add_run(' (December 2014) appears to have lapsed \u2014 brownfield register entry ended July 2021')

add_body('')
add_para('Is that the right site?', italic=True)

add_separator()

# --- SECTION 2: Six Refusal Reasons ---
add_heading_styled('Our Understanding of the Six Refusal Reasons', level=2)

add_body(
    'Based on the policy framework and your description, here\u2019s what our system identified as the '
    'likely reasons. I\u2019d be interested to know how close we are:'
)

# Reason 1
add_heading_styled('1. Poor Design in the Conservation Area', level=3)
add_body(
    'Cornwall Local Plan Policy 24 (Historic Environment) requires development to sustain the cultural '
    'distinctiveness and significance of the historic environment. The Conservation Area Appraisal for '
    'Camborne Town Centre would set the character parameters \u2014 scale, massing, materials, roofscape. '
    'Our read is the original design didn\u2019t respond to this context adequately.'
)

# Reason 2
add_heading_styled('2. Poor Design in the World Heritage Site', level=3)
add_body(
    'The WHS Supplementary Planning Document (2017) requires applicants to demonstrate that the '
    'Outstanding Universal Value of the mining landscape is preserved. Camborne is specifically '
    'highlighted as an area where "spatial arrangements are of particular concern and may be vulnerable '
    'unless planning policies and guidance are rigorously and consistently applied." We assume the '
    'original scheme didn\u2019t include a WHS impact assessment or demonstrate how the design related '
    'to the mining heritage character.'
)

# Reason 3
add_heading_styled('3. Inadequate BNG Assessment', level=3)
add_body(
    'Mandatory BNG came into force 12 February 2024 for major development (10+ dwellings). Cornwall '
    'Council had actually been requesting BNG information on major applications locally since March 2020. '
    'A 10-apartment scheme on a 0.05ha brownfield site is classified as major \u2014 our assumption is '
    'the original application either omitted the BNG metric entirely or submitted it without the full '
    'statement, draft Habitat Management and Monitoring Plan, or draft S106 terms that Cornwall requires '
    'as a local validation requirement for majors.'
)

# Reason 4
add_heading_styled('4. Lack of Energy Assessment', level=3)
add_body(
    'Cornwall\u2019s Climate Emergency DPD was adopted February 2023, with Policy SEC1 (Sustainable '
    'Energy and Construction) in force from 15 June 2023. This is a validation requirement. If the '
    'original application was submitted without an energy statement addressing SEC1, it\u2019s a '
    'straightforward refusal reason \u2014 the document simply wasn\u2019t there.'
)

# Reason 5
add_heading_styled('5. Section 106 Contributions', level=3)
add_body(
    'Here\u2019s where the designated area status really bites. Because the site sits within both a '
    'Conservation Area and the WHS, it\u2019s a designated protected area \u2014 which means Cornwall\u2019s '
    'affordable housing threshold drops to more than 5 dwellings (not the standard 10). For schemes of '
    '6\u201310 dwellings in designated areas, a financial contribution is sought per unit of affordable '
    'housing. Our assumption is the original applicant didn\u2019t offer any S106 heads of terms at all '
    '\u2014 possibly unaware the designated area threshold is lower.'
)

# Reason 6
add_heading_styled('6. Sixth Reason (Our Best Guess)', level=3)
add_body(
    'Possibly related to scale/massing impact on the streetscape, overlooking/amenity for neighbouring '
    'properties, or inadequate parking/servicing for the site\u2019s constraints. On a 0.05ha plot, '
    '10 apartments would be tight. Would be interested to know what this one was.'
)

add_separator()

# --- SECTION 3: Strategic Move ---
add_heading_styled('Why the 10 to 9 Reduction Was the Key Strategic Move', level=2)

add_body(
    'This is the part that impressed us most. On the surface, dropping one unit looks like a minor '
    'concession. In reality, it fundamentally changes the planning classification:'
)

# Table
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Shading Accent 1'

headers = ['', '10 Units', '9 Units']
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(10)

data = [
    ['Development class', 'Major', 'Minor'],
    ['BNG requirements', 'Full statement + metric + draft HMMP + draft S106', 'National minimum info only'],
    ['Determination period', '13 weeks', '8 weeks'],
    ['Committee likelihood', 'Higher \u2014 major apps attract more scrutiny', 'More likely officer delegated'],
    ['Affordable housing', 'Required (>5 in designated area)', 'Still required \u2014 financial contribution'],
]

for idx, row_data in enumerate(data):
    add_table_row_data(table, idx + 1, row_data, bold_first=True)

doc.add_paragraph('')  # spacer

add_body(
    'The affordable housing obligation doesn\u2019t change (both above the >5 designated area threshold), '
    'but the overall regulatory burden drops significantly. Combined with a complete redesign that properly '
    'responds to the Conservation Area character and WHS context, plus the missing documents now included '
    '(BNG, Energy Statement, S106 heads of terms), you\u2019ve essentially rebuilt the application from '
    'scratch while also reducing the procedural complexity.'
)

add_body(
    'The decision to resubmit rather than appeal was correct too. Our analysis of appeal outcomes shows '
    'that refusals with 3+ substantive reasons (and yours had 6, including missing documents) have very '
    'low appeal success rates. A fresh application with a new team was always going to be the better path.'
)

add_separator()

# --- SECTION 4: System Demo ---
add_heading_styled('What Our System Would Flag at Site Analysis Stage', level=2)

add_body(
    'Just to give you a sense of what we\u2019re building \u2014 when we ran the postcode through our '
    'site analysis pipeline, it automatically detected:'
)

add_bullet('Conservation Area + WHS dual designation (flagged as high heritage sensitivity)')
add_bullet('Brownfield status with lapsed permission')
add_bullet('Designated area affordable housing threshold (>5 dwellings, not >10)')
add_bullet('23 constraint API checks completed (flood, ecology, heritage, geology, contamination)')
add_bullet(
    'Auto-selected reports: Heritage Impact Assessment, BNG Screening, Energy Statement, '
    'Design & Access Statement, Planning Statement, CIL/S106 Assessment \u2014 all as mandatory '
    'before any submission'
)
add_bullet('Cross-report triggers: 8 heritage-related consistency checks across all documents')
add_bullet(
    'Consultee prediction: WHS Team, Conservation Officer, Town Council all flagged for early engagement'
)

add_body('')
add_body(
    'The idea is that a system like this, used at the outset, would have caught every one of those '
    'six refusal reasons before the first application was ever submitted.'
)

add_separator()

# --- SECTION 5: Questions ---
add_heading_styled('Over to You', level=2)

add_body(
    'I genuinely think this is a great case study in planning consultancy done right \u2014 not just '
    'submitting documents, but fundamentally rethinking the strategy. The fact you secured full Town '
    'Council and officer support on the resubmission speaks to getting the pre-app engagement right '
    'from the start.'
)

add_body('A few questions if you have a moment:')

doc.add_paragraph('Have we got the right site?', style='List Number')
doc.add_paragraph('How close are we on the six refusal reasons?', style='List Number')
doc.add_paragraph(
    'Was the major-to-minor threshold shift a deliberate part of the strategy from the outset?',
    style='List Number'
)
doc.add_paragraph(
    'Did you use the WHS pre-application advice service, and was it useful?',
    style='List Number'
)
doc.add_paragraph(
    'What was the new design team\u2019s approach to the Conservation Area character \u2014 '
    'contemporary or traditional?',
    style='List Number'
)

add_body('')
add_body(
    'No pressure on any of this \u2014 genuinely just learning. But if you\u2019re ever interested in '
    'seeing what our system does in more detail, happy to run a site through it for you.'
)

add_body('Great work, Dan. Well deserved result.')

# Sign off
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Best regards,')
p = doc.add_paragraph()
run = p.add_run('Chris')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('PF & Co Site Intelligence')
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

# --- Save ---
output_path = os.path.join(
    r'C:\Users\chris\pf-&-co-construction  - web\docs',
    'Email-Dan-Mitchell-Camborne-Case-Study.docx'
)
doc.save(output_path)
print(f'Document saved to: {output_path}')
