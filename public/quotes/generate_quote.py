from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

AMBER = RGBColor(0xF5, 0x9E, 0x0B)
BLACK = RGBColor(0x0A, 0x0A, 0x0A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0x99, 0x99, 0x99)
MID_GREY = RGBColor(0xF5, 0xF5, 0xF5)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_run(paragraph, text, bold=False, size=None, color=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = 'Calibri'
    return run

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

# ========================================
# HEADER
# ========================================
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
header_table.columns[0].width = Cm(10)
header_table.columns[1].width = Cm(7)

# Left: company name
left_cell = header_table.cell(0, 0)
p = left_cell.paragraphs[0]
add_run(p, 'PF & Co Construction Ltd', bold=True, size=18, color=BLACK)
p2 = left_cell.add_paragraph()
add_run(p2, 'Structural Engineering & Construction — Surrey, London & the South East', size=8, color=GREY)
p2.paragraph_format.space_before = Pt(2)

# Right: contact
right_cell = header_table.cell(0, 1)
p = right_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, 'PF & Co Construction Ltd\n', bold=True, size=9, color=BLACK)
add_run(p, '01483 363 210\n', size=9, color=GREY)
add_run(p, 'info@pfcoconstruction.co.uk\n', size=9, color=GREY)
add_run(p, 'www.pfcoconstruction.co.uk', size=9, color=GREY)

# Amber divider
divider = doc.add_paragraph()
divider.paragraph_format.space_before = Pt(6)
divider.paragraph_format.space_after = Pt(6)
run = divider.add_run()
# Add a border line using a thin table
div_table = doc.add_table(rows=1, cols=1)
div_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = div_table.cell(0, 0)
cell.text = ''
set_cell_shading(cell, 'F59E0B')
set_row_height(div_table.rows[0], 0.1)
# Remove the empty paragraph before the divider table
divider._element.getparent().remove(divider._element)

# ========================================
# QUOTATION TITLE BAR
# ========================================
title_table = doc.add_table(rows=1, cols=2)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
title_table.columns[0].width = Cm(10)
title_table.columns[1].width = Cm(7)

left = title_table.cell(0, 0)
set_cell_shading(left, '0A0A0A')
p = left.paragraphs[0]
add_run(p, 'QUOTATION', bold=True, size=14, color=WHITE)

right = title_table.cell(0, 1)
set_cell_shading(right, '0A0A0A')
p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, 'Ref: PF/Q/2026-0302/WBC  |  Date: 2 March 2026', size=9, color=AMBER)

set_row_height(title_table.rows[0], 0.9)

doc.add_paragraph()  # spacer

# ========================================
# CLIENT & SITE INFO
# ========================================
info_table = doc.add_table(rows=1, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.columns[0].width = Cm(8.5)
info_table.columns[1].width = Cm(8.5)

# Client
client_cell = info_table.cell(0, 0)
p = client_cell.paragraphs[0]
add_run(p, 'CLIENT', bold=True, size=8, color=LIGHT_GREY)
p2 = client_cell.add_paragraph()
add_run(p2, 'Wonersh Bowls Club\n', bold=True, size=10, color=BLACK)
add_run(p2, 'Wonersh Common\nWonersh, Guildford\nSurrey GU5 0PH', size=10, color=GREY)

# Site
site_cell = info_table.cell(0, 1)
p = site_cell.paragraphs[0]
add_run(p, 'SITE & PROJECT', bold=True, size=8, color=LIGHT_GREY)
p2 = site_cell.add_paragraph()
add_run(p2, 'Construction of Disabled Parking Area\n', bold=True, size=10, color=BLACK)
add_run(p2, 'Wonersh Bowls Club, Wonersh Common, Surrey\n', size=10, color=GREY)
add_run(p2, 'Area: 40 m²  |  Duration: approx. 1 week\n', size=10, color=GREY)
add_run(p2, 'Site electrics, water & welfare provided by client', size=8, color=LIGHT_GREY, italic=True)

# Add borders to info cells
for cell in [client_cell, site_cell]:
    set_cell_border(cell,
        top={"val": "single", "sz": "4", "color": "E0E0E0"},
        bottom={"val": "single", "sz": "4", "color": "E0E0E0"},
        left={"val": "single", "sz": "4", "color": "E0E0E0"},
        right={"val": "single", "sz": "4", "color": "E0E0E0"})

doc.add_paragraph()  # spacer

# ========================================
# SCOPE OF WORKS
# ========================================
scope_p = doc.add_paragraph()
scope_p.paragraph_format.space_after = Pt(4)
add_run(scope_p, 'Scope of Works', bold=True, size=11, color=BLACK)

scope_body = doc.add_paragraph()
scope_body.paragraph_format.space_after = Pt(10)
add_run(scope_body,
    'Formation of a new disabled-accessible parking area (40 m²) comprising excavation to formation level, '
    'MOT Type 1 sub-base installation (150 mm compacted depth), heavy-duty weed prevention membrane, '
    'commercial-grade gravel stabilisation grid system suitable for vehicular loading and DDA-compliant '
    'wheelchair access, gravel infill, and 18 linear metres of granite set kerb edging. Works include careful '
    'excavation around existing underground services (soil pipes, electrical cables and water supply pipes) '
    'and full disposal of excavated material.',
    size=10, color=GREY)

# ========================================
# SITE LOCATION
# ========================================
loc_p = doc.add_paragraph()
loc_p.paragraph_format.space_after = Pt(4)
add_run(loc_p, 'Site Location', bold=True, size=11, color=BLACK)

loc_body = doc.add_paragraph()
loc_body.paragraph_format.space_after = Pt(2)
add_run(loc_body, 'Wonersh Common, Wonersh, Guildford, Surrey GU5 0PH', size=10, color=GREY)

loc_link = doc.add_paragraph()
loc_link.paragraph_format.space_after = Pt(10)
add_run(loc_link, 'Google Maps: ', size=9, color=GREY)
add_run(loc_link, 'https://maps.google.com/?q=Wonersh+Bowls+Club+Surrey+GU5+0PH', size=9, color=RGBColor(0x00, 0x56, 0xB3))

# ========================================
# PRICING TABLE
# ========================================

pricing_data = [
    # (ref, description, amount, is_section_header, is_subtotal, is_total, is_vat, is_grand)
    ('1.0', 'Excavation & Site Clearance', '', True, False, False, False, False),
    ('1.1', 'Excavate existing ground over 40 m² area to formation level (approx. 200 mm depth). Strip existing turf and topsoil. Work carefully around existing underground services including soil pipes, electrical cables and water supply pipes. Hand-dig where required in proximity to services.', '850.00', False, False, False, False, False),
    ('1.2', 'Disposal of excavated material (soil/turf) via grab lorry — 1 no. load, clean inert waste', '440.00', False, False, False, False, False),
    ('1.3', '1.5 tonne mini excavator hire — 1 week including delivery to site and collection on completion', '560.00', False, False, False, False, False),
    ('', 'Section 1 Subtotal', '1,850.00', False, True, False, False, False),

    ('2.0', 'Type 1 Sub-base', '', True, False, False, False, False),
    ('2.1', 'Supply and deliver MOT Type 1 crushed stone aggregate to site — 10 tonne load', '500.00', False, False, False, False, False),
    ('2.2', 'Spread and compact Type 1 sub-base to 150 mm finished compacted depth over 40 m² area, including level checks', '420.00', False, False, False, False, False),
    ('', 'Section 2 Subtotal', '920.00', False, True, False, False, False),

    ('3.0', 'Weed Prevention Membrane', '', True, False, False, False, False),
    ('3.1', 'Supply and install heavy-duty geotextile weed prevention membrane over compacted sub-base. Minimum 150 mm overlaps at all joints and edges. Approx. 45 m² including overlaps.', '225.00', False, False, False, False, False),
    ('', 'Section 3 Subtotal', '225.00', False, True, False, False, False),

    ('4.0', 'Gravel Stabilisation Grid System', '', True, False, False, False, False),
    ('4.1', 'Supply and install commercial heavy-duty gravel stabilisation grid panels to 40 m² area. Grids rated for vehicular loading and DDA-compliant for disabled wheelchair access (min. 250 t/m² load rating). Cut and fit to perimeter and around any service access points.', '1,520.00', False, False, False, False, False),
    ('', 'Section 4 Subtotal', '1,520.00', False, True, False, False, False),

    ('5.0', 'Gravel Infill', '', True, False, False, False, False),
    ('5.1', 'Supply, deliver and install approximately 3 tonnes of 20 mm angular gravel to fill grid cells. Levelled flush with top of grid panels and lightly compacted.', '410.00', False, False, False, False, False),
    ('', 'Section 5 Subtotal', '410.00', False, True, False, False, False),

    ('6.0', 'Granite Kerb Edging', '', True, False, False, False, False),
    ('6.1', 'Install 18 linear metres of granite set kerb edging to perimeter of disabled parking area. Bedded on semi-dry sand/cement mortar with concrete haunch to rear. Line and level throughout. Granite sets supplied by client. All bedding and haunching materials (cement, building sand, concrete) supplied by contractor.', '775.00', False, False, False, False, False),
    ('', 'Section 6 Subtotal', '775.00', False, True, False, False, False),

    # Totals
    ('', 'Total (excluding VAT)', '£5,700.00', False, False, True, False, False),
    ('', 'VAT @ 20%', '£1,140.00', False, False, False, True, False),
    ('', 'Total (including VAT)', '£6,840.00', False, False, False, False, True),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.columns[0].width = Cm(1.2)
table.columns[1].width = Cm(13.0)
table.columns[2].width = Cm(2.8)

# Header row
for i, text in enumerate(['Ref', 'Description', 'Amount (£)']):
    cell = table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i == 2 else WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, text, bold=True, size=8, color=WHITE)
    set_cell_shading(cell, '0A0A0A')

# Data rows
for ref, desc, amount, is_header, is_sub, is_total, is_vat, is_grand in pricing_data:
    row = table.add_row()

    if is_header:
        for i, text in enumerate([ref, desc, amount]):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, bold=True, size=10, color=BLACK)
            set_cell_shading(cell, 'F5F5F5')
    elif is_sub:
        for i, text in enumerate([ref, desc, amount]):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, bold=True, size=9, color=BLACK)
            set_cell_shading(cell, 'FAFAFA')
            set_cell_border(cell, top={"val": "single", "sz": "4", "color": "CCCCCC"},
                           bottom={"val": "single", "sz": "8", "color": "E0E0E0"})
    elif is_total:
        for i, text in enumerate([ref, desc, amount]):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_run(p, text, bold=True, size=12, color=AMBER)
            else:
                add_run(p, text, bold=True, size=11, color=WHITE)
            set_cell_shading(cell, '0A0A0A')
    elif is_vat:
        for i, text in enumerate([ref, desc, amount]):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, bold=True, size=9, color=BLACK)
            set_cell_shading(cell, 'F5F5F5')
    elif is_grand:
        for i, text in enumerate([ref, desc, amount]):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, text, bold=True, size=11, color=BLACK)
            set_cell_shading(cell, 'F59E0B')
    else:
        # Normal row
        row.cells[0].paragraphs[0].text = ''
        add_run(row.cells[0].paragraphs[0], ref, bold=True, size=9, color=GREY)
        row.cells[1].paragraphs[0].text = ''
        add_run(row.cells[1].paragraphs[0], desc, size=9, color=RGBColor(0x33, 0x33, 0x33))
        row.cells[2].paragraphs[0].text = ''
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, amount, size=9, color=BLACK)

doc.add_paragraph()  # spacer

# ========================================
# NOTES: INCLUDED / EXCLUDED
# ========================================
notes_heading = doc.add_paragraph()
notes_heading.paragraph_format.space_after = Pt(6)
add_run(notes_heading, 'Important Notes', bold=True, size=11, color=BLACK)

notes_table = doc.add_table(rows=1, cols=2)
notes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
notes_table.columns[0].width = Cm(8.5)
notes_table.columns[1].width = Cm(8.5)

# Included
inc_cell = notes_table.cell(0, 0)
p = inc_cell.paragraphs[0]
add_run(p, 'INCLUDED IN THIS QUOTATION', bold=True, size=8, color=LIGHT_GREY)

included_items = [
    'All labour, materials, plant and disposal as described',
    '1.5 tonne mini excavator for duration of works (1 week)',
    'Grab lorry for muck away (1 load)',
    '10 tonnes MOT Type 1 aggregate delivered',
    'All bedding materials for granite kerb installation',
    'Commercial-grade gravel grids for vehicular and disabled access',
    'Approx. 3 tonnes of 20 mm gravel infill',
]
for item in included_items:
    p = inc_cell.add_paragraph()
    add_run(p, f'•  {item}', size=8.5, color=GREY)
    p.paragraph_format.space_after = Pt(1)

# Excluded
exc_cell = notes_table.cell(0, 1)
p = exc_cell.paragraphs[0]
add_run(p, 'EXCLUDED / BY OTHERS', bold=True, size=8, color=LIGHT_GREY)

excluded_items = [
    'Site electrics, water supply and welfare facilities (by client)',
    'Granite kerb sets (provided by client, on site)',
    'Disabled parking signage, road markings and bay delineation',
    'Any additional drainage works not described',
    'Service diversions or protection works beyond hand-digging',
    'Additional loads of Type 1 if 10 tonnes proves insufficient',
    'Any works to surfaces beyond the 40 m² area',
]
for item in excluded_items:
    p = exc_cell.add_paragraph()
    add_run(p, f'•  {item}', size=8.5, color=GREY)
    p.paragraph_format.space_after = Pt(1)

doc.add_paragraph()

# ========================================
# ASSUMPTIONS
# ========================================
assume_heading = doc.add_paragraph()
assume_heading.paragraph_format.space_after = Pt(6)
add_run(assume_heading, 'Assumptions & Advisory Notes', bold=True, size=11, color=BLACK)

notes = [
    ('Type 1 Quantity: ',
     '10 tonnes of Type 1 at a compacted density of approx. 1.9 t/m³ will yield approximately 5.3 m³ loose, '
     'compacting to around 4.2 m³ or roughly 105 mm depth over 40 m². To achieve the full 150 mm compacted depth, '
     'approximately 14 tonnes may be required. We have priced a single 10-tonne load as instructed; an additional '
     'delivery can be arranged at cost if required on site.'),
    ('Underground Services: ',
     'We have allowed for careful hand-digging around known services. Should any unforeseen services, obstructions '
     'or ground conditions (e.g. rock, contaminated soil, high water table) be encountered, additional costs may '
     'apply and will be agreed before proceeding.'),
    ('Access: ',
     'Quotation assumes adequate vehicle access for a grab lorry and aggregate delivery vehicle to reach the works area.'),
]
for title, body in notes:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, size=9, color=BLACK)
    add_run(p, body, size=9, color=GREY)

doc.add_paragraph()

# ========================================
# TERMS & CONDITIONS
# ========================================
terms_heading = doc.add_paragraph()
terms_heading.paragraph_format.space_after = Pt(6)
add_run(terms_heading, 'Terms & Conditions', bold=True, size=11, color=BLACK)

terms = [
    'This quotation is valid for 30 days from the date shown above.',
    'All prices are in pounds sterling (£ GBP). VAT is charged at the prevailing rate (currently 20%).',
    'Payment terms: 50% deposit on acceptance of quotation; balance due on completion of works.',
    'The estimated programme is approximately 5 working days, subject to weather and site conditions.',
    'Any variations to the described scope of works will be agreed in writing before commencement and may result in additional costs.',
    'PF & Co Construction Ltd holds full public liability insurance and employer\'s liability insurance. Certificates available on request.',
    'Works will be carried out in accordance with all relevant Building Regulations, British Standards and CDM Regulations where applicable.',
    'The client is responsible for obtaining any necessary planning permissions or consents prior to commencement.',
    'We reserve the right to sub-contract specialist elements of the works where necessary.',
    'This quotation is subject to the full terms and conditions of PF & Co Construction Ltd, available on request.',
]
for i, term in enumerate(terms, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f'{i}.  ', bold=True, size=8.5, color=GREY)
    add_run(p, term, size=8.5, color=GREY)

doc.add_paragraph()

# ========================================
# ACCEPTANCE
# ========================================
accept_heading = doc.add_paragraph()
accept_heading.paragraph_format.space_after = Pt(6)
add_run(accept_heading, 'Acceptance', bold=True, size=11, color=BLACK)

accept_body = doc.add_paragraph()
accept_body.paragraph_format.space_after = Pt(16)
add_run(accept_body,
    'I/we accept this quotation and authorise PF & Co Construction Ltd to proceed with the works as described above.',
    size=9, color=GREY)

# Signature lines
sig_table = doc.add_table(rows=2, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sig_table.columns[0].width = Cm(8.5)
sig_table.columns[1].width = Cm(8.5)

sig_labels = [
    ('Signed (Client)', 'Position / Authority'),
    ('Print Name', 'Date'),
]
for row_i, (left_label, right_label) in enumerate(sig_labels):
    for col_i, label in enumerate([left_label, right_label]):
        cell = sig_table.cell(row_i, col_i)
        # Add spacing
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(24)
        add_run(p, '', size=9)
        set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": "CCCCCC"})
        p2 = cell.add_paragraph()
        add_run(p2, label, size=8, color=LIGHT_GREY)

doc.add_paragraph()

# ========================================
# FOOTER
# ========================================
# Amber divider
div_table2 = doc.add_table(rows=1, cols=1)
div_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = div_table2.cell(0, 0)
cell.text = ''
set_cell_shading(cell, 'F59E0B')
set_row_height(div_table2.rows[0], 0.1)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(6)
add_run(footer_p, 'PF & Co Construction Ltd', bold=True, size=8, color=GREY)
add_run(footer_p, '  |  01483 363 210  |  info@pfcoconstruction.co.uk  |  www.pfcoconstruction.co.uk', size=8, color=LIGHT_GREY)
footer_p2 = doc.add_paragraph()
footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer_p2, 'Registered in England & Wales  |  © 2020–2026 PF & Co Construction Ltd. All Rights Reserved.', size=7, color=LIGHT_GREY)

# ========================================
# SAVE
# ========================================
output_dir = os.path.dirname(os.path.abspath(__file__))
docx_path = os.path.join(output_dir, 'PF-Quote-Wonersh-Bowls-Club-Disabled-Parking.docx')
doc.save(docx_path)
print(f'Word document saved: {docx_path}')
