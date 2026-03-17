"""Generate PF & Co Subscription Model Research as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

h1_style = doc.styles['Heading 1']
h1_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
h1_style.font.size = Pt(18)

h2_style = doc.styles['Heading 2']
h2_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
h2_style.font.size = Pt(14)

h3_style = doc.styles['Heading 3']
h3_style.font.color.rgb = RGBColor(0x2D, 0x5F, 0x8A)
h3_style.font.size = Pt(12)

ACCENT = RGBColor(0x1B, 0x3A, 0x5C)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, '1B3A5C')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()
    return table


def bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PF & Co Site Intelligence')
run.font.size = Pt(28)
run.font.color.rgb = ACCENT
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Subscription Model Research\n& Revenue Projections')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2D, 0x5F, 0x8A)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Date: 13 March 2026\n').font.size = Pt(11)
meta.add_run('Status: Draft for review\n').font.size = Pt(11)
meta.add_run('Classification: Internal').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. COMPETITIVE LANDSCAPE
# ============================================================
doc.add_heading('1. Competitive Landscape', level=1)

doc.add_heading('Direct Competitors & Pricing', level=2)

add_table(
    ['Platform', 'What They Do', 'Price', 'Key Limitation'],
    [
        ['LandInsight (LandTech)', 'Site sourcing + planning data', '\u00a360\u2013\u00a3180/month', 'Data only \u2014 no reports'],
        ['Nimbus Maps', 'Data layers + mapping', '\u00a3150\u2013\u00a3390/month', 'Data only \u2014 no reports'],
        ['SearchLand', 'Ownership + constraints mapping', '\u00a3195+/month', 'Data only \u2014 no reports'],
        ['PlanningBot', 'AI site appraisals', '\u00a349/month', '15 data sources, 1 report type, not submission-ready'],
        ['Groundsure', 'Environmental screening', '\u00a350\u2013\u00a3200/report', 'Per-report, no planning context'],
        ['Landmark', 'Environmental data products', '\u00a340\u2013\u00a3200/report', 'Raw data, no recommendations'],
        ['Glenigan', 'Construction pipeline leads', '\u00a38K\u2013\u00a350K/year', 'Enterprise only, no site-level reports'],
        ['Barbour ABI', 'Construction pipeline leads', 'Custom (similar to Glenigan)', 'Enterprise only, no site-level reports'],
        ['PropertyData', 'Market analytics', '\u00a314\u2013\u00a360/month', 'Analytics only, no planning reports'],
    ],
)

doc.add_heading('Detailed Competitor Breakdown', level=2)

doc.add_heading('LandInsight (LandTech)', level=3)
doc.add_paragraph('Starter: \u00a360/month \u2014 25 planning app views, 100 ownership lookups')
doc.add_paragraph('Pro: \u00a3180/month \u2014 300 ownership lookups, unlimited planning views')
doc.add_paragraph('Unlimited: Custom pricing \u2014 advanced policy/environmental/demographic insights')
doc.add_paragraph('Also offers LandEnhance (planning application management) at custom pricing. Claims to save 4\u20135 hours per project.')

doc.add_heading('Nimbus Maps', level=3)
doc.add_paragraph('Elite: \u00a3150/month (\u00a3125/month annual = \u00a31,500/year)')
doc.add_paragraph('Elite+: \u00a3270/month (\u00a3225/month annual = \u00a32,700/year)')
doc.add_paragraph('Elite MAX: \u00a3390/month (\u00a3325/month annual = \u00a33,900/year)')
doc.add_paragraph('All tiers include unlimited title lookups and OS MasterMap credits. Land Registry title deeds: \u00a37 per document on top.')

doc.add_heading('SearchLand', level=3)
doc.add_paragraph('Standard: \u00a3195/month per user (billed annually) \u2014 unlimited title lookups, 500 OS MasterMap credits, planning applications')
doc.add_paragraph('Pro: Custom pricing \u2014 adds strategic land tools, BNG tools, energy/DNO data, Scotland data')
doc.add_paragraph('Claims 20% cheaper than LandInsight equivalent.')

doc.add_heading('PlanningBot \u2014 Closest Direct Competitor', level=3)
doc.add_paragraph('Free: \u00a30/month \u2014 5 site assessments, basic flood & heritage checks')
doc.add_paragraph('Professional: \u00a349/month \u2014 unlimited site assessments, 15+ data sources, 9 tools, DOCX/PDF reports')
doc.add_paragraph('Generates site appraisal reports, interactive constraint maps, authority profiles, appeal decision research (90K+ decisions). However: only 1 report type, not submission-ready, 15 data sources only.')

doc.add_heading('Glenigan', level=3)
doc.add_paragraph('Custom pricing only; estimated \u00a31,000\u2013\u00a350,000/year depending on scope. Typical mid-market subscription: \u00a38,000\u2013\u00a315,000/year. Includes 5 user logins, monthly intelligence reports, account management.')

doc.add_page_break()

# ============================================================
# 2. MARKET GAP
# ============================================================
doc.add_heading('2. The Market Gap PF & Co Fills', level=1)

doc.add_paragraph(
    'No other platform generates actual submission-ready, multi-domain technical reports. '
    'The existing market splits into:'
)
doc.add_paragraph('Data platforms (LandInsight, Nimbus, SearchLand) \u2014 show data, leave you to write reports', style='List Bullet')
doc.add_paragraph('Environmental data brokers (Groundsure, Landmark) \u2014 sell raw data per-search', style='List Bullet')
doc.add_paragraph('Traditional consultants \u2014 \u00a31,000\u20135,000+ per report, 4\u20138 weeks, multiple firms', style='List Bullet')
doc.add_paragraph('Nascent AI tools (PlanningBot) \u2014 basic site appraisals, not submission-ready', style='List Bullet')

bold_para('PF & Co is the only platform producing 45 document types from 58 data sources with QA pipeline, case law integration, and cross-report validation.')

add_table(
    ['PF & Co Advantage', 'Detail'],
    [
        ['Breadth', '58 data sources, 112+ agents, 45 document types'],
        ['Depth', 'Submission-ready DOCX/PDF with regulatory citations, 170 case law precedents'],
        ['Speed', '48-hour turnaround vs 4\u20138 week traditional'],
        ['Coverage', 'Planning + environmental + engineering + financial in one pack'],
        ['Price disruption', 'Complete Intelligence at \u00a36,995 vs \u00a315,000\u2013\u00a350,000+ traditional'],
    ],
)

doc.add_page_break()

# ============================================================
# 3. TRADITIONAL CONSULTANT PRICING
# ============================================================
doc.add_heading('3. Traditional Consultant Pricing Comparison', level=1)

add_table(
    ['Report Type', 'Traditional Fee', 'PF & Co EA Price', 'Saving'],
    [
        ['Flood Risk Assessment (desktop)', '\u00a31,000\u2013\u00a32,500', '\u00a3375', '63\u201385%'],
        ['Preliminary Ecological Appraisal', '\u00a3348\u2013\u00a31,200', '\u00a3495 (BNG)', '~50%'],
        ['BNG Assessment', '\u00a31,500\u2013\u00a37,500', '\u00a3495', '67\u201393%'],
        ['Heritage Impact Assessment', '\u00a3500\u2013\u00a32,000', '\u00a3545', '~50%'],
        ['Noise Impact Assessment', '\u00a31,000\u2013\u00a32,000', '\u00a3445', '55\u201378%'],
        ['Air Quality Assessment', '\u00a31,200\u2013\u00a31,350', '\u00a3395', '67\u201371%'],
        ['Transport Assessment', '\u00a31,500\u2013\u00a35,000', '\u00a3495', '67\u201390%'],
        ['Arboricultural Survey (BS5837)', '\u00a3300\u2013\u00a31,000', '\u00a3575', '~40%'],
        ['Phase 1 Contamination Desk Study', '\u00a3500\u2013\u00a31,200', '\u00a3595', '~40%'],
        ['Daylight/Sunlight Assessment', '\u00a32,000+', '\u00a3595', '70%+'],
        ['Design & Access Statement', '\u00a3875\u2013\u00a32,000', '\u00a3395', '55\u201380%'],
        ['Planning Statement', '\u00a31,000\u2013\u00a33,000', '\u00a3495', '50\u201384%'],
        ['Energy Statement', '\u00a3800\u2013\u00a32,000', '\u00a3445', '44\u201378%'],
    ],
)

bold_para('Full planning application pack (10\u201350 dwellings):')
doc.add_paragraph('Traditional total professional fees: \u00a315,000\u2013\u00a350,000+', style='List Bullet')
doc.add_paragraph('PF & Co Complete Intelligence: \u00a36,995', style='List Bullet')
doc.add_paragraph('Saving: 55\u201386%', style='List Bullet')

doc.add_page_break()

# ============================================================
# 4. FULL REPORT CAPABILITY
# ============================================================
doc.add_heading('4. Full Report Capability (45 Document Types)', level=1)

doc.add_heading('Purchasable Reports \u2014 24 on Website', level=2)

for section, items in [
    ('Core Intelligence', [
        'Site Acquisition Intelligence (SAI)', 'Development Finance Summary (DFS)',
        'Site Feasibility Report (SFR)', 'Geotechnical Desk Study (GDS)',
        'Flood Risk Assessment (FRA)', 'Phase 1 Contamination Assessment (P1C)',
    ]),
    ('Planning & Strategy', [
        'Planning Statement (PS)', 'Pre-Application Advice Pack (PAA)',
        'Design & Access Statement (DAS)', 'Concept Feasibility Study (CFS)',
        'CIL Liability Assessment (CIL)',
    ]),
    ('Environmental & Ecology', [
        'Biodiversity Net Gain Screening (BNG)', 'Energy Statement (ES)',
        'Air Quality Screening (AQS)',
    ]),
    ('Specialist Assessments', [
        'Heritage Impact Assessment (HIA)', 'Transport Statement (TS)',
        'Parking Survey (PKS)', 'Arboricultural Desk Study (ADS)',
        'Noise Impact Assessment (NIA)', 'Daylight & Sunlight Assessment (DSA)',
    ]),
    ('Construction Readiness', [
        'Construction Management Plan (CMP)', 'Design Readiness Review (DRR)',
    ]),
    ('Compliance & Legal', [
        'Building Control (BC)', 'Party Wall Assessment (PWA)',
    ]),
]:
    bold_para(section)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Constraint-Triggered Reports \u2014 15 Additional', level=2)
for item in [
    'Affordable Housing Statement', 'Drainage Strategy', 'Housing Needs Assessment',
    'Utilities Capacity Assessment', 'Open Space Strategy', 'Travel Plan',
    'Community Involvement Statement', 'EIA Screening', 'Socio-Economic Assessment',
    'Phasing & Delivery Plan', 'Landscape & Visual Impact Assessment (LVIA)',
    'Wind Microclimate Screening', 'Fire Statement', 'Archaeological Assessment',
    'Arboricultural Impact Assessment',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Pack Documents \u2014 6', level=2)
for item in [
    'Application Covering Letter', 'Application Summary', 'Submission Pack Index',
    'LPA Validation Checklist', 'Risk Summary Sheet', 'Client Decision Pack',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 5. DEMAND SIGNALS
# ============================================================
doc.add_heading('5. Market Demand Signals', level=1)

add_table(
    ['Indicator', 'Figure', 'Source'],
    [
        ['UK PropTech market (2024)', '\u00a33.32 billion', 'Market Research Future'],
        ['UK PropTech projected (2035)', '\u00a325.08 billion (20% CAGR)', 'Market Research Future'],
        ['UK PropTech revenue by 2030', '\u00a36.56 billion (14.3% CAGR)', 'Grand View Research'],
        ['UK SaaS market by 2026', '\u00a315.4 billion (nearly double)', 'Industry reports'],
        ['Construction tech funding (2024)', '\u00a34.8 billion (32% of PropTech)', 'Qubit Capital'],
        ['SME construction software CAGR', '11.06% (fastest growing)', 'Fortune Business Insights'],
        ['Cloud construction mgmt CAGR', '13.00%', 'Fortune Business Insights'],
        ['Company software already SaaS', '80% (expected 85% by 2026)', 'InventoryBase'],
        ['MHCLG PropTech Challenge interest', '116 expressions, 84 applications', 'MHCLG Digital'],
        ['PropTech time reduction claims', '80\u201390% via automated assessments', 'Qubit Capital'],
    ],
)

doc.add_page_break()

# ============================================================
# 6. PROPOSED SUBSCRIPTION TIERS
# ============================================================
doc.add_heading('6. Proposed Subscription Tiers', level=1)

doc.add_heading('Operating Model: Manual Verification', level=2)
doc.add_paragraph(
    'No self-serve interface currently \u2014 clients submit sites, PF & Co runs the system and verifies outputs. '
    'Every project requires human QA time. Tiers must have hard project caps reflecting delivery capacity.'
)

doc.add_heading('Tier Structure', level=2)

add_table(
    ['', 'Scout', 'Professional', 'Enterprise'],
    [
        ['Monthly fee', '\u00a3199/month', '\u00a3499/month', '\u00a3999/month'],
        ['Annual fee', '\u00a31,990/year (save \u00a3398)', '\u00a34,990/year (save \u00a3998)', '\u00a39,990/year (save \u00a31,998)'],
        ['Screenings included', '2/month', '4/month', '8/month'],
        ['Full report packs included', '0', '1/month', '2/month'],
        ['Turnaround', '5 working days', '48\u201372 hours', '48 hours priority'],
        ['Discount on extra reports', '15%', '25%', '35%'],
        ['Support', 'Email', 'Email + phone', 'Dedicated contact'],
        ['Target', 'Architects, small developers', 'Active developers, planning consultants', 'Volume housebuilders, land promoters'],
    ],
)

doc.add_heading('Value Anchoring', level=2)
doc.add_paragraph('Scout at \u00a3199/month sits between PlanningBot (\u00a349) and LandInsight Pro (\u00a3180) \u2014 but delivers actual verified reports', style='List Bullet')
doc.add_paragraph('Professional at \u00a3499/month undercuts a single traditional consultant report (\u00a31,000\u20132,500) while giving 4 screenings + a full pack', style='List Bullet')
doc.add_paragraph('Enterprise at \u00a3999/month is a fraction of Glenigan (\u00a38K\u201350K/year) and includes actual deliverables', style='List Bullet')
doc.add_paragraph('A Professional subscriber gets roughly \u00a310K\u2013\u00a315K worth of traditional consultant work per month', style='List Bullet')

doc.add_page_break()

# ============================================================
# 7. REVENUE PROJECTIONS
# ============================================================
doc.add_heading('7. Revenue Projections', level=1)

doc.add_heading('Scenario 1: Conservative (Month 3\u20136)', level=2)
doc.add_paragraph('Getting started, word of mouth, LinkedIn traction')

add_table(
    ['Tier', 'Subscribers', 'Monthly Revenue'],
    [
        ['Scout', '5', '\u00a3995'],
        ['Professional', '2', '\u00a3998'],
        ['Enterprise', '0', '\u00a30'],
        ['Subscription total', '7', '\u00a31,993'],
        ['Extra report purchases (est.)', '\u2014', '\u00a32,000'],
        ['MONTHLY TOTAL', '', '\u00a33,993'],
        ['ANNUAL RUN RATE', '', '\u00a347,916'],
    ],
)
doc.add_paragraph('Workload: ~18 screenings + 2 full packs = ~70\u201380 hours/month. Manageable solo.')

doc.add_heading('Scenario 2: Moderate (Month 6\u201312)', level=2)
doc.add_paragraph('Established reputation, repeat clients, referrals flowing')

add_table(
    ['Tier', 'Subscribers', 'Monthly Revenue'],
    [
        ['Scout', '12', '\u00a32,388'],
        ['Professional', '5', '\u00a32,495'],
        ['Enterprise', '1', '\u00a3999'],
        ['Subscription total', '18', '\u00a35,882'],
        ['Extra report purchases (est.)', '\u2014', '\u00a34,500'],
        ['MONTHLY TOTAL', '', '\u00a310,382'],
        ['ANNUAL RUN RATE', '', '\u00a3124,584'],
    ],
)
doc.add_paragraph('Workload: ~44 screenings + 7 full packs = ~150\u2013160 hours/month. At capacity solo \u2014 time to hire.')

doc.add_heading('Scenario 3: Growth (Month 12\u201318)', level=2)
doc.add_paragraph('Team of 2\u20133, system proven, marketing engine running')

add_table(
    ['Tier', 'Subscribers', 'Monthly Revenue'],
    [
        ['Scout', '20', '\u00a33,980'],
        ['Professional', '10', '\u00a34,990'],
        ['Enterprise', '3', '\u00a32,997'],
        ['Subscription total', '33', '\u00a311,967'],
        ['Extra report purchases (est.)', '\u2014', '\u00a38,000'],
        ['MONTHLY TOTAL', '', '\u00a319,967'],
        ['ANNUAL RUN RATE', '', '\u00a3239,604'],
    ],
)
doc.add_paragraph('Workload: ~80 screenings + 16 full packs = ~350 hours/month. Needs 2\u20132.5 FTE.')

doc.add_page_break()

# ============================================================
# 8. CAPACITY PLANNING
# ============================================================
doc.add_heading('8. Capacity Planning', level=1)

doc.add_heading('Time Per Deliverable', level=2)
add_table(
    ['Deliverable', 'System Time', 'Human QA/Verification', 'Total'],
    [
        ['Site screening', '~1 hour (automated)', '2\u20133 hours', '3\u20134 hours'],
        ['Full report pack', '~2\u20134 hours (automated)', '8\u201312 hours', '10\u201316 hours'],
    ],
)

doc.add_heading('Revenue Ceiling by Team Size', level=2)
add_table(
    ['Team Size', 'Screenings/month', 'Full Packs/month', 'Revenue Ceiling'],
    [
        ['1 person (you)', '~20', '~8', '~\u00a310K/month'],
        ['You + 1 hire', '~40', '~16', '~\u00a320K/month'],
        ['Team of 3', '~60', '~24', '~\u00a335K/month'],
    ],
)

doc.add_heading('Key Insight', level=2)
doc.add_paragraph(
    'The bottleneck is human verification, not system capacity. As QA experience builds, '
    'verification time will reduce (3 hours to 1.5 hours per screening). Pattern recognition improves \u2014 '
    'common site types become faster. QA checklists can be templated for junior reviewers.'
)

doc.add_page_break()

# ============================================================
# 9. EXISTING BUNDLES
# ============================================================
doc.add_heading('9. Existing Bundle Pricing (Reference)', level=1)

add_table(
    ['Bundle', 'EA Price', 'RRP', 'Reports Included'],
    [
        ['The Complete Intelligence', '\u00a36,995', '\u00a319,020', '22 reports'],
        ['The Full Planning Suite', '\u00a33,495', '\u00a37,285', '7 reports'],
        ['The Developer Due Diligence', '\u00a33,095', '\u00a36,415', '7 reports'],
        ['The Appeal-Ready Pack', '\u00a32,635', '\u00a34,395', '4 reports'],
        ['The Self-Build Starter', '\u00a33,050', '\u00a35,590', '4 reports'],
        ['The Architect Support Pack', '\u00a32,625', '\u00a34,725', '4 reports'],
        ['The Triple Threat', '\u00a32,375', '\u00a34,600', '3 reports'],
        ['The Construction Readiness Pack', '\u00a31,195', '\u00a32,480', '2 reports'],
    ],
)

doc.add_page_break()

# ============================================================
# 10. RECOMMENDATIONS
# ============================================================
doc.add_heading('10. Recommendations', level=1)

recommendations = [
    ('Launch with Scout + Professional only',
     'Don\'t offer Enterprise until the delivery model is proven and you have confidence in turnaround times at volume.'),
    ('Cap total subscribers initially (e.g. 15)',
     'Scarcity creates urgency and protects quality. "Founding member" positioning builds loyalty and case studies.'),
    ('Track hours per project type from day one',
     'This data tells you when to hire, what to reprice, and where the system needs to get faster.'),
    ('The upsell is where the real money is',
     'Subscriptions get clients in the door. Discounted full report packs drive the majority of revenue.'),
    ('Consider a "Founding Member" discount',
     'First 10 subscribers get 20% off for life \u2014 builds loyalty, generates case studies, and creates advocates.'),
    ('Review pricing quarterly',
     'As verification gets faster, margins improve and tier limits can increase. Don\'t lock yourself into year-one pricing forever.'),
]

for i, (title, detail) in enumerate(recommendations, 1):
    bold_para(f'{i}. {title}')
    doc.add_paragraph(detail)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('This document is a working draft for internal planning purposes.')
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = r'C:\Users\chris\pf-&-co-construction  - web\docs\PF-Co-Subscription-Model-Research.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
